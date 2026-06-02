from docx import Document

doc = Document('now.docx')

print("Word Document Content:")
print("=" * 50)
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n\nTables in Document:")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text, end=" | ")
        print()
